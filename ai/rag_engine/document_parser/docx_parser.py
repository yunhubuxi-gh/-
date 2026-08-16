"""
DOCX（Word）文档解析器（python-docx）

解析策略（与 PDF 解析器完全对齐，走统一三通道）：
1. 按文档顺序遍历段落与表格，提取纯文本（正文）
2. 解析 docx 的 zip 包与 rels 关系，完整提取段落内、表格内的全部内嵌图片二进制
3. 每一张图片调用 PaddleOCR 识别文字，OCR 结果并入文本一起参与统一分块
4. 返回 ParsedDocument：
   - pages：按「逻辑页」组织的文本（正文 + 每张图片的 OCR 文本），供文本分块器统一分片
   - images：全部内嵌图片（内存态），供上层 document_service 落盘 + 多模态向量化
5. 容错铁则：单张图片提取失败 / OCR 失败 / 图片损坏，只跳过该图、记录警告，
   绝不中断整篇 docx 解析；正文文本与其余图片照常入库。

说明：图片的多模态向量化由上层 rag_pipeline.ingest_images 统一处理（逐张 try-except），
本解析器只负责「文本通道（正文 + 图片 OCR）」与「图片二进制提取」。
"""
from __future__ import annotations

import os
from typing import List, Tuple, Iterator, Union, Optional, Callable

from config.settings import settings
from ai.rag_engine.document_parser.base_parser import (
    BaseDocumentParser,
    ParsedDocument,
    PageText,
    ExtractedImage,
)
from utils.exceptions import FileOperationException
from utils.error_codes import DOC_PARSE_FAILED
from utils.text_utils import clean_text
from utils.logger import get_logger

logger = get_logger(__name__)

# 近似一页的字符数（Word 无真实分页，按内容量切分逻辑页，与旧版保持一致）
CHARS_PER_APPROX_PAGE = 1500

# OpenXML 命名空间
_NS = {
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "v": "urn:schemas-microsoft-com:vml",
    "o": "urn:schemas-microsoft-com:office:office",
}
_R_EMBED = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
_O_ID = "{urn:schemas-microsoft-com:office:office}id"


class DocxParser(BaseDocumentParser):
    supported_extensions = [".docx"]

    def __init__(self, ocr_engine=None):
        """
        Args:
            ocr_engine: OCR 引擎实例，缺省时懒加载 utils.ocr_engine.get_ocr_engine()
        """
        self._ocr_engine = ocr_engine

    def _get_ocr_engine(self):
        if self._ocr_engine is None:
            from utils.ocr_engine import get_ocr_engine
            self._ocr_engine = get_ocr_engine()
        return self._ocr_engine

    def parse(
        self,
        file_path: str,
        progress_callback: Optional[Callable[[int, int], None]] = None,
    ) -> ParsedDocument:
        """
        Args:
            file_path: docx 文件路径
            progress_callback: 可选进度回调 `(done, total)`，每完成一张图片 OCR 调用一次，
                用于上报细粒度子阶段进度（前端展示「OCR识别中 x/总数」）。
                回调可能在并发 OCR 的多个工作线程中触发，实现需自行保证线程安全。
        """
        if not os.path.exists(file_path):
            raise FileOperationException(DOC_PARSE_FAILED, f"文件不存在: {file_path}")

        try:
            from docx import Document
        except ImportError as e:
            raise FileOperationException(
                DOC_PARSE_FAILED, "python-docx 未安装，请执行 pip install python-docx"
            ) from e

        try:
            doc = Document(file_path)
        except Exception as e:
            raise FileOperationException(DOC_PARSE_FAILED, f"DOCX 打开失败: {e}") from e

        # 顺序扫描正文 + 图片，产出 pages 与 images（图片绑定逻辑页，OCR 文本参与分块）
        pages, images = self._scan_to_pages_and_images(doc, progress_callback=progress_callback)

        title = os.path.splitext(os.path.basename(file_path))[0]
        logger.info(
            f"DOCX 解析完成: 提取到 {len(images)} 张内嵌图片, "
            f"pages={len(pages)}, chars={sum(len(p.text) for p in pages)}"
        )
        return ParsedDocument(
            title=title,
            pages=pages,
            images=images,
            metadata={
                "source": file_path,
                "format": "docx",
                "page_count": len(pages),
            },
        )

    # ---------- 顺序扫描：正文 + 图片 → 逻辑页 ----------

    def _scan_to_pages_and_images(
        self,
        doc,
        progress_callback: Optional[Callable[[int, int], None]] = None,
    ) -> Tuple[List[PageText], List[ExtractedImage]]:
        """按文档顺序遍历段落与表格，产出逻辑页文本（正文+图片OCR）与内嵌图片列表。

        顺序扫描规则（与旧版 _split_into_pages 的「逐段累积，超 1500 字归档」一致）：
        - 正文段落 / 表格行 → 计入当前逻辑页文本
        - 图片 → 绑定当前逻辑页号（页内序号从 0 递增），其 OCR 文本同样计入当前逻辑页
        - 当前页累计字符达到阈值 → 归档该页，开启新页

        性能增强：先收集全部内容单元，再对图片并发 OCR（ThreadPoolExecutor），
        最后按文档顺序组装逻辑页——OCR 文本与正文的分页/分块语义与串行版完全一致。
        """
        # 1. 收集全部内容单元（保持文档顺序）
        units = list(self._iter_ordered_units(doc))

        # 2. 收集所有图片并并发 OCR，结果按原顺序回填
        image_positions: List[int] = []   # 图片在原 units 中的下标
        all_images: List[ExtractedImage] = []
        for idx, (kind, payload) in enumerate(units):
            if kind == "image":
                image_positions.append(idx)
                all_images.append(payload)
        ocr_texts = self._ocr_images_concurrent(all_images, progress_callback=progress_callback)
        ocr_map = {pos: ocr_texts[j] for j, pos in enumerate(image_positions)}

        # 3. 按文档顺序组装逻辑页（与串行版一致的 flush/add 语义）
        pages: List[PageText] = []
        images: List[ExtractedImage] = []

        cur_texts: List[str] = []
        cur_len = 0
        cur_page_images: List[ExtractedImage] = []

        def flush_page() -> None:
            nonlocal cur_len
            page_no = len(pages) + 1
            for idx, img in enumerate(cur_page_images):
                img.page_number = page_no
                img.index = idx
            pages.append(PageText(page_number=page_no, text="\n".join(cur_texts)))
            images.extend(cur_page_images)
            cur_texts.clear()
            cur_page_images.clear()
            cur_len = 0

        def add_text(text: str) -> None:
            nonlocal cur_len
            if not text:
                return
            cur_texts.append(text)
            cur_len += len(text)

        for idx, (kind, payload) in enumerate(units):
            if kind == "text":
                add_text(payload)
            else:  # image
                img: ExtractedImage = payload
                cur_page_images.append(img)
                add_text(ocr_map.get(idx, ""))

            if cur_len >= CHARS_PER_APPROX_PAGE:
                flush_page()

        # 收尾：剩余内容（或空文档兜底 1 个空页）归档
        if cur_texts or cur_page_images or not pages:
            flush_page()

        return pages, images

    def _ocr_images_concurrent(
        self,
        images: List[ExtractedImage],
        progress_callback: Optional[Callable[[int, int], None]] = None,
    ) -> List[str]:
        """并发 OCR 一批图片，返回与输入同序的文本列表（单张失败仅该张为空串）。

        说明：PaddleOCR 底层 C++ 预测器非线程安全，多线程并发 predict 会段错误，
        故并发走**进程池**（utils.ocr_pool，每个子进程独立 PaddleOCR 实例）；
        进程池不可用时自动降级为串行（ocr_engine 内部锁保护 predict）。
        """
        from utils.ocr_pool import ocr_images_concurrent

        if not images:
            return []
        image_bytes_list = [img.image_bytes for img in images]
        try:
            concurrency = int(getattr(settings, "ocr_concurrency", 4) or 4)
        except (TypeError, ValueError):
            concurrency = 4
        return ocr_images_concurrent(
            image_bytes_list,
            progress_callback=progress_callback,
            max_workers=concurrency,
        )

    def _iter_ordered_units(self, doc) -> Iterator[Tuple[str, Union[str, ExtractedImage]]]:
        """按文档顺序产出内容单元：(kind, payload)。

        kind="text"  → payload 为文本字符串
        kind="image" → payload 为 ExtractedImage（已含二进制与扩展名，page_number 待后续绑定）
        """
        try:
            from docx.table import Table
            from docx.text.paragraph import Paragraph
        except ImportError:
            # 极老版本兜底：仅段落
            for para in doc.paragraphs:
                text = clean_text(para.text)
                if text:
                    yield ("text", text)
            return

        seen_rids = set()

        for child in doc.element.body.iterchildren():
            tag = child.tag
            if tag.endswith("}p"):
                para = Paragraph(child, doc)
                text = clean_text(para.text)
                if text:
                    yield ("text", text)
                # 段落内图片（含嵌套文本框等）
                for img in self._extract_images_from_element(child, doc, seen_rids):
                    yield ("image", img)
            elif tag.endswith("}tbl"):
                table = Table(child, doc)
                # 表格内文本
                for row in table.rows:
                    row_text = clean_text(" | ".join(c.text.strip() for c in row.cells))
                    if row_text:
                        yield ("text", row_text)
                # 表格内图片（递归查找所有单元格）
                for img in self._extract_images_from_element(child, doc, seen_rids):
                    yield ("image", img)

    # ---------- 图片提取（zip + rels 关系）----------

    def _extract_images_from_element(self, element, doc, seen_rids: set) -> Iterator[ExtractedImage]:
        """从 XML 元素（段落 / 表格）递归提取内嵌图片，按文档顺序产出，去重。

        支持两种图片存储方式：
        - 现代 DrawingML：a:blip 的 r:embed 属性指向图片关系 id
        - 老式 VML：v:imagedata 的 o:id 属性指向图片关系 id
        """
        rids: List[str] = []
        for blip in element.findall(".//a:blip", _NS):
            rid = blip.get(_R_EMBED)
            if rid:
                rids.append(rid)
        for imagedata in element.findall(".//v:imagedata", _NS):
            rid = imagedata.get(_O_ID)
            if rid:
                rids.append(rid)

        for rid in rids:
            if rid in seen_rids:
                continue
            seen_rids.add(rid)
            img = self._resolve_image(rid, doc)
            if img is not None:
                yield img

    def _resolve_image(self, rid: str, doc) -> ExtractedImage:
        """通过关系 id 从 docx 包内解析图片二进制与扩展名，失败返回 None（跳过该图）"""
        try:
            rel = doc.part.rels[rid]
        except KeyError:
            logger.warning(f"docx 图片关系不存在，跳过: rId={rid}")
            return None
        try:
            image_part = rel.target_part
            blob = image_part.blob
        except Exception as e:
            logger.warning(f"docx 图片二进制读取失败，跳过: rId={rid}, err={e}")
            return None
        if not blob:
            return None
        ext = self._resolve_ext(image_part)
        return ExtractedImage(page_number=0, image_bytes=blob, ext=ext, index=0)

    @staticmethod
    def _resolve_ext(image_part) -> str:
        """解析图片扩展名（优先 content_type，回退 partname，兜底 png）"""
        ct = getattr(image_part, "content_type", "") or ""
        if ct.startswith("image/"):
            ext = ct.split("/")[-1].lower()
            if ext == "jpeg":
                return "jpg"
            if ext in ("png", "jpg", "gif", "bmp", "webp", "tiff"):
                return ext
        try:
            ext = image_part.partname.ext.lstrip(".").lower()
        except Exception:
            ext = ""
        if ext == "jpeg":
            return "jpg"
        return ext if ext in ("png", "jpg", "gif", "bmp", "webp", "tiff") else "png"

    # ---------- 图片 OCR ----------

    def _ocr_image(self, image_bytes: bytes) -> str:
        """对单张图片做 OCR，失败返回空字符串（优雅降级，不影响图片向量化）"""
        engine = self._get_ocr_engine()
        if engine is None:
            return ""
        try:
            return clean_text(engine.recognize_image_bytes(image_bytes) or "")
        except Exception as e:
            logger.warning(f"docx 图片 OCR 失败（跳过该图 OCR，不影响整体）: {e}")
            return ""
